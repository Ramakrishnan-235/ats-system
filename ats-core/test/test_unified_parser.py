import pytest
import io
import fitz
import docx
from PIL import Image, ImageDraw
from ats_core.parsers.unified_parser import UnifiedDocumentParser

def create_sample_docx_bytes() -> bytes:
    doc = docx.Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Principal Distributed Systems Engineer • jane.doe@example.com")
    doc.add_heading("Core Technical Skills", level=2)
    doc.add_paragraph("Python, FastAPI, Kubernetes, PostgreSQL, Redis, Docker, AWS")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Cloud Systems Inc — 2021 to Present")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def create_sample_resume_image() -> bytes:
    img = Image.new("RGB", (1200, 1600), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((80, 80), "Alex Rivera", fill=(10, 10, 10))
    draw.text((80, 130), "Staff Site Reliability Engineer", fill=(60, 60, 60))
    draw.text((80, 190), "Skills: Python, Kubernetes, Docker, AWS", fill=(20, 20, 20))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def create_sample_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((50, 72), "Dr. Marcus Vance\nStaff Architect\nPython, FastAPI, AWS", fontsize=12)
    return doc.write()

def test_unified_parser_detects_formats():
    parser = UnifiedDocumentParser()
    
    pdf_bytes = create_sample_pdf_bytes()
    assert parser.detect_format(pdf_bytes, "resume.pdf") == "pdf"

    docx_bytes = create_sample_docx_bytes()
    assert parser.detect_format(docx_bytes, "resume.docx") == "docx"

    img_bytes = create_sample_resume_image()
    assert parser.detect_format(img_bytes, "resume.png") == "image"

def test_unified_parser_parses_docx():
    parser = UnifiedDocumentParser()
    docx_bytes = create_sample_docx_bytes()
    text, engine, fmt = parser.parse(docx_bytes, "resume.docx")

    assert fmt == "docx"
    assert "Jane Doe" in text
    assert engine == "python-docx"

def test_unified_parser_parses_image():
    parser = UnifiedDocumentParser()
    img_bytes = create_sample_resume_image()
    text, engine, fmt = parser.parse(img_bytes, "resume.png")

    assert fmt == "image"
    assert len(text.strip()) > 30

def test_unified_parser_parses_pdf():
    parser = UnifiedDocumentParser()
    pdf_bytes = create_sample_pdf_bytes()
    text, engine, fmt = parser.parse(pdf_bytes, "resume.pdf")

    assert fmt == "pdf"
    assert "Marcus Vance" in text
