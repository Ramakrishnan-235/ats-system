import io
import pytest
import docx
import fitz
from PIL import Image, ImageDraw

from ats_core.parsers.unified_parser import UnifiedDocumentParser
from ats_core.parsers.resume_parser import parse_resume_to_candidate, extract_text_from_document
from ats_core.parsers.docx_parser import DocxResumeParser
from ats_core.parsers.image_ocr_parser import ImageOCRResumeParser


def generate_test_pdf() -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    text = (
        "Sarah Connor\n"
        "Lead Site Reliability Engineer\n"
        "sarah.c@example.com • (415) 555-0199 • San Francisco, CA\n"
        "linkedin.com/in/sarahconnor\n\n"
        "CORE TECHNICAL SKILLS\n"
        "Python, Kubernetes, Docker, AWS, Terraform, Prometheus, Go, PostgreSQL\n\n"
        "PROFESSIONAL EXPERIENCE\n"
        "Lead SRE — Cloudflare (2020 — Present)\n"
        "Spearheaded enterprise infrastructure and Kubernetes cluster reliability.\n\n"
        "EDUCATION\n"
        "Bachelor of Science in Computer Science — UC Berkeley\n"
    )
    page.insert_text((50, 72), text, fontsize=11)
    return doc.write()


def generate_test_docx() -> bytes:
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "David Chen • david.chen@example.com • (212) 555-0144 • New York, NY"

    doc.add_heading("David Chen", level=1)
    doc.add_paragraph("Principal Platform Architect")

    doc.add_heading("Technical Proficiencies", level=2)
    doc.add_paragraph("Python, FastAPI, Kubernetes, PostgreSQL, Redis, Docker, Microservices, Git")

    doc.add_heading("Professional Experience", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Role & Organization"
    hdr_cells[1].text = "Timeline"
    hdr_cells[2].text = "Impact"

    r1 = table.add_row().cells
    r1[0].text = "Staff Architect — Robinhood"
    r1[1].text = "2021 — Present"
    r1[2].text = "Engineered real-time trading backend microservices with sub-millisecond p99 latency."

    doc.add_heading("Education", level=2)
    doc.add_paragraph("M.S. in Computer Science — Carnegie Mellon University")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_test_image() -> bytes:
    img = Image.new("RGB", (1400, 1800), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    draw.text((60, 60), "Elena Vance", fill=(10, 10, 10))
    draw.text((60, 110), "Senior AI Platform Dev • elena.v@example.com • (555) 789-0123 • Seattle, WA", fill=(50, 50, 50))
    draw.text((60, 170), "Core Skills: Python, PyTorch, FastAPI, Docker, Kubernetes, LangChain, PostgreSQL", fill=(20, 20, 20))
    draw.text((60, 240), "Work Experience: Senior AI Engineer at Anthropic (2021 — Present)", fill=(20, 20, 20))
    draw.text((60, 300), "Scaled LLM inference serving clusters and distributed training workflows.", fill=(40, 40, 40))
    draw.text((60, 370), "Education: B.S. in Computer Engineering — University of Washington", fill=(20, 20, 20))

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def test_unified_parser_multi_format():
    parser = UnifiedDocumentParser()

    # 1. Test PDF
    pdf_bytes = generate_test_pdf()
    assert parser.detect_format(pdf_bytes, "sarah.pdf") == "pdf"
    pdf_text, pdf_eng, pdf_fmt = parser.parse(pdf_bytes, "sarah.pdf")
    assert pdf_fmt == "pdf"
    assert "Sarah Connor" in pdf_text
    assert "Kubernetes" in pdf_text

    # 2. Test DOCX
    docx_bytes = generate_test_docx()
    assert parser.detect_format(docx_bytes, "david.docx") == "docx"
    docx_text, docx_eng, docx_fmt = parser.parse(docx_bytes, "david.docx")
    assert docx_fmt == "docx"
    assert docx_eng == "python-docx"
    assert "David Chen" in docx_text
    assert "FastAPI" in docx_text
    assert "|" in docx_text  # Table extracted

    # 3. Test Image OCR
    img_bytes = generate_test_image()
    assert parser.detect_format(img_bytes, "elena.png") == "image"
    img_text, img_eng, img_fmt = parser.parse(img_bytes, "elena.png")
    assert img_fmt == "image"
    assert len(img_text.strip()) > 40


def test_parse_resume_to_candidate_pdf():
    pdf_bytes = generate_test_pdf()
    candidate = parse_resume_to_candidate(
        pdf_bytes,
        filename="sarah_connor.pdf",
        target_job={"title": "Staff Infrastructure SRE", "department": "Platform"}
    )

    assert candidate["name"] == "Sarah Connor"
    assert "sarah.c@example.com" in candidate["email"]
    assert "Kubernetes" in candidate["core_skills"]
    assert "Python" in candidate["core_skills"]
    assert candidate["years_of_experience"] >= 4.0
    assert candidate["scorecard"]["overall_match_score"] >= 80


def test_parse_resume_to_candidate_docx():
    docx_bytes = generate_test_docx()
    candidate = parse_resume_to_candidate(
        docx_bytes,
        filename="david_chen.docx",
        target_job={"title": "Principal Platform Architect", "department": "Core Systems"}
    )

    assert "David Chen" in candidate["name"] or "david.chen" in candidate["email"]
    assert "FastAPI" in candidate["core_skills"]
    assert "PostgreSQL" in candidate["core_skills"]
    assert candidate["years_of_experience"] >= 3.0
    assert candidate["scorecard"]["overall_match_score"] >= 80


def test_parse_resume_to_candidate_image():
    img_bytes = generate_test_image()
    candidate = parse_resume_to_candidate(
        img_bytes,
        filename="elena_vance_resume.png",
        target_job={"title": "Senior AI Platform Dev", "department": "Machine Learning"}
    )

    assert candidate["name"] is not None
    assert len(candidate["core_skills"]) >= 3
    assert candidate["scorecard"]["overall_match_score"] >= 75
    assert len(candidate["scorecard"]["categories"]) >= 2
