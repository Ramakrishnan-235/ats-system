import io
import pytest
import docx
from ats_core.parsers.docx_parser import DocxResumeParser

def create_sample_docx_bytes() -> bytes:
    doc = docx.Document()
    
    # Add Header
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Jane Doe • jane.doe@example.com • (555) 123-4567 • San Francisco, CA"

    # Add Title / Name
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Principal Distributed Systems Engineer")

    # Add Core Skills
    doc.add_heading("Core Technical Skills", level=2)
    doc.add_paragraph("Python, FastAPI, Kubernetes, PostgreSQL, Redis, Docker, AWS, Kafka")

    # Add Experience Table
    doc.add_heading("Work Experience", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Company & Role"
    hdr_cells[1].text = "Period"
    hdr_cells[2].text = "Key Achievements"

    row1 = table.add_row().cells
    row1[0].text = "Cloud Platform Inc — Staff Architect"
    row1[1].text = "2021 — Present"
    row1[2].text = "Architected real-time microservices handling 50k RPS with zero downtime."

    # Add Education
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.S. in Computer Science — Stanford University")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def test_docx_parser_extracts_content_and_tables():
    docx_bytes = create_sample_docx_bytes()
    parser = DocxResumeParser()
    extracted_text, engine = parser.parse_docx(docx_bytes, filename="jane_doe.docx")

    assert engine == "python-docx"
    assert "Jane Doe" in extracted_text
    assert "jane.doe@example.com" in extracted_text
    assert "Kubernetes" in extracted_text
    assert "Cloud Platform Inc" in extracted_text
    assert "50k RPS" in extracted_text
    assert "Stanford University" in extracted_text
    # Verify markdown table conversion
    assert "|" in extracted_text
    assert "---" in extracted_text

def test_docx_parser_empty_stream_raises():
    parser = DocxResumeParser()
    with pytest.raises(ValueError, match="empty"):
        parser.parse_docx(b"", filename="empty.docx")
